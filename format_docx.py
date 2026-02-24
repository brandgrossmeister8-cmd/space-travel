#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт форматирования Word документа ФИНАЛЬНЫЙ_ОТЧЁТ_ЗАЩИТА.docx
Шрифт: Tahoma, Цвет: голубой, Таблицы на всю ширину
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

# Голубой цвет (основной)
BLUE_MAIN = RGBColor(0, 102, 204)      # #0066CC - голубой
BLUE_LIGHT = "CCE5FF"                   # Светло-голубой для фона
BLUE_HEADER = "0066CC"                  # Голубой для заголовков таблиц

# Ширина страницы A4: 21 см, поля 1.5 см с каждой стороны = 18 см рабочая область
PAGE_WIDTH_CM = 18
PAGE_WIDTH_TWIPS = int(PAGE_WIDTH_CM * 567)  # 10206 twips

def set_cell_shading(cell, color):
    """Устанавливает цвет фона ячейки"""
    tcPr = cell._tc.get_or_add_tcPr()
    # Удаляем старый shading если есть
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def set_cell_width(cell, width_twips):
    """Устанавливает ширину ячейки в twips"""
    tcPr = cell._tc.get_or_add_tcPr()
    # Удаляем старую ширину
    for tcW in tcPr.findall(qn('w:tcW')):
        tcPr.remove(tcW)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)

def format_table_full_width(table, total_width_twips):
    """Форматирует таблицу на полную ширину с равномерными столбцами"""
    tbl = table._tbl

    # Считаем количество столбцов из первой строки
    num_cols = 0
    if len(table.rows) > 0:
        num_cols = len(table.rows[0].cells)

    if num_cols == 0:
        return

    col_width = total_width_twips // num_cols

    # Получаем или создаём tblPr
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Удаляем старые настройки
    for tag in ['w:tblW', 'w:tblLayout', 'w:tblBorders', 'w:jc']:
        for elem in tblPr.findall(qn(tag)):
            tblPr.remove(elem)

    # Ширина таблицы
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.insert(0, tblW)

    # Фиксированная раскладка
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # Границы таблицы - голубые
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), BLUE_HEADER)
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Удаляем старый tblGrid
    for old_tblGrid in tbl.findall(qn('w:tblGrid')):
        tbl.remove(old_tblGrid)

    # Создаём новый tblGrid с правильными ширинами
    tblGrid = OxmlElement('w:tblGrid')
    for i in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(col_width))
        tblGrid.append(gridCol)

    # Вставляем tblGrid после tblPr
    tblPr_index = list(tbl).index(tblPr)
    tbl.insert(tblPr_index + 1, tblGrid)

    # Устанавливаем ширину каждой ячейки
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, col_width)

def format_document(input_path, output_path):
    """Основная функция форматирования"""
    print("Загрузка документа...")
    doc = Document(input_path)

    # 1. Установка размеров страницы A4 и полей 1.5 см
    print("Настройка полей страницы (1.5 см)...")
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 2. Настройка стилей - шрифт Tahoma
    print("Настройка стилей (Tahoma, голубой цвет)...")
    styles = doc.styles

    # Стиль Normal
    style_normal = styles['Normal']
    style_normal.font.name = 'Tahoma'
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.space_after = Pt(3)
    style_normal.paragraph_format.line_spacing = 1.0

    # Стили заголовков
    for i in range(1, 4):
        try:
            style = styles[f'Heading {i}']
            style.font.name = 'Tahoma'
            style.font.bold = True
            style.font.color.rgb = BLUE_MAIN
            if i == 1:
                style.font.size = Pt(14)
                style.paragraph_format.space_before = Pt(8)
                style.paragraph_format.space_after = Pt(4)
            elif i == 2:
                style.font.size = Pt(12)
                style.paragraph_format.space_before = Pt(6)
                style.paragraph_format.space_after = Pt(3)
            else:
                style.font.size = Pt(11)
                style.paragraph_format.space_before = Pt(4)
                style.paragraph_format.space_after = Pt(2)
        except:
            pass

    # 3. Удаление ВСЕХ пустых параграфов и форматирование текста
    print("Удаление пустых параграфов и форматирование текста...")
    paragraphs_to_remove = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if text == '':
            paragraphs_to_remove.append(para)
        else:
            # Минимальные отступы
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0

            # Шрифт Tahoma для всех runs
            for run in para.runs:
                run.font.name = 'Tahoma'
                if run.font.size is None:
                    run.font.size = Pt(10)

            # Форматирование заголовков - голубой цвет
            if text.startswith('# ') or text.startswith('РАЗДЕЛ') or text.startswith('ПРИЛОЖЕНИЕ'):
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(4)
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(14)
                    run.font.color.rgb = BLUE_MAIN
            elif text.startswith('## '):
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = BLUE_MAIN
            elif text.startswith('### '):
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = BLUE_MAIN

    # Удаление пустых параграфов
    print(f"Удаление {len(paragraphs_to_remove)} пустых параграфов...")
    for para in reversed(paragraphs_to_remove):
        p = para._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    # 4. Форматирование таблиц - на всю ширину
    print(f"Форматирование таблиц ({len(doc.tables)} шт.) на ширину {PAGE_WIDTH_CM} см...")
    for table in doc.tables:
        # Устанавливаем полную ширину таблицы
        format_table_full_width(table, PAGE_WIDTH_TWIPS)

        # Форматирование ячеек
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Форматирование текста в ячейках
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.font.name = 'Tahoma'
                        run.font.size = Pt(9)

                # Заголовок таблицы (первая строка) - голубой фон
                if row_idx == 0:
                    set_cell_shading(cell, BLUE_HEADER)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                # Чередование цветов строк
                elif row_idx % 2 == 0:
                    set_cell_shading(cell, BLUE_LIGHT)

    # 5. Сохранение
    print(f"Сохранение в {output_path}...")
    doc.save(output_path)
    print("Готово!")

if __name__ == "__main__":
    input_file = "/Users/Rita/Documents/Vibecoding/14_Space Travel _отработка_Курсор/ФИНАЛЬНЫЙ_ОТЧЁТ_ЗАЩИТА.docx"
    output_file = "/Users/Rita/Documents/Vibecoding/14_Space Travel _отработка_Курсор/ФИНАЛЬНЫЙ_ОТЧЁТ_ЗАЩИТА_FORMATTED.docx"
    format_document(input_file, output_file)
