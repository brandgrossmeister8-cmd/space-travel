#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт добавления диаграмм в Word документ
"""

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Без GUI
plt.rcParams['font.family'] = 'Arial'
plt.rcParams['font.size'] = 10

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Папка для временных изображений
TEMP_DIR = "/Users/Rita/Documents/Vibecoding/14_Space Travel _отработка_Курсор/temp_charts"
os.makedirs(TEMP_DIR, exist_ok=True)

def create_pie_chart_departments(filename):
    """Диаграмма: Распределение сотрудников по категориям"""
    labels = ['Генератор выручки\n132 чел. (67%)', 'Усилитель выручки\n31 чел. (16%)',
              'Обеспечивающая\n30 чел. (15%)', 'Без отдела\n5 чел. (2%)']
    sizes = [132, 31, 30, 5]
    colors = ['#003366', '#0066CC', '#3399FF', '#99CCFF']
    explode = (0.02, 0, 0, 0)

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='',
           shadow=False, startangle=90)
    ax.axis('equal')
    plt.title('Распределение 198 сотрудников по категориям департаментов', fontsize=11, fontweight='bold')
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_bar_chart_automation(filename):
    """Диаграмма: Экономия по волнам автоматизации"""
    waves = ['Волна 1\n(быстрые победы)', 'Волна 2\n(системные)', 'Волна 3\n(стратегические)']
    savings_hours = [400, 500, 200]
    savings_money = [2.9, 6.0, 4.0]

    fig, ax1 = plt.subplots(figsize=(7, 4))

    x = range(len(waves))
    width = 0.35

    bars1 = ax1.bar([i - width/2 for i in x], savings_hours, width, label='Экономия (ч/мес)', color='#003366')
    ax1.set_ylabel('Часов в месяц', color='#003366')
    ax1.tick_params(axis='y', labelcolor='#003366')
    ax1.set_xticks(x)
    ax1.set_xticklabels(waves)

    ax2 = ax1.twinx()
    bars2 = ax2.bar([i + width/2 for i in x], savings_money, width, label='Экономия (млн ₽/год)', color='#FF6600')
    ax2.set_ylabel('Млн рублей в год', color='#FF6600')
    ax2.tick_params(axis='y', labelcolor='#FF6600')

    plt.title('Экономия по волнам автоматизации', fontsize=11, fontweight='bold')
    fig.legend(loc='upper right', bbox_to_anchor=(0.88, 0.88))
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_pie_chart_automation_categories(filename):
    """Диаграмма: Распределение точек автоматизации по категориям"""
    labels = ['Рутина\n35 точек (55%)', 'Узкое место\n16 точек (25%)', 'Интеграция\n13 точек (20%)']
    sizes = [35, 16, 13]
    colors = ['#003366', '#FF6600', '#339933']
    explode = (0.02, 0, 0)

    fig, ax = plt.subplots(figsize=(5, 4))
    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='',
           shadow=False, startangle=90)
    ax.axis('equal')
    plt.title('64 точки автоматизации по категориям', fontsize=11, fontweight='bold')
    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_horizontal_bar_tz(filename):
    """Диаграмма: Экономия по ТЗ"""
    tz_names = ['ТЗ-01: САМО↔Битрикс', 'ТЗ-02: Шаблоны', 'ТЗ-07: Power BI',
                'ТЗ-04: Маршрутизация', 'ТЗ-12: GDS→1С', 'ТЗ-13: Динам.цены']
    savings = [90, 90, 86, 84, 81, 55]

    fig, ax = plt.subplots(figsize=(7, 4))
    colors = ['#003366', '#003366', '#003366', '#0066CC', '#FF6600', '#FF6600']
    bars = ax.barh(tz_names, savings, color=colors)
    ax.set_xlabel('Экономия (ч/мес)')
    ax.set_title('ТОП-6 ТЗ по экономии времени', fontsize=11, fontweight='bold')

    for bar, val in zip(bars, savings):
        ax.text(val + 2, bar.get_y() + bar.get_height()/2, f'{val} ч', va='center', fontsize=9)

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_timeline_chart(filename):
    """Диаграмма: Дорожная карта внедрения"""
    fig, ax = plt.subplots(figsize=(8, 3))

    # Волны
    waves = [
        ('Подготовка', 0, 1, '#99CCFF'),
        ('Волна 1: 8 ТЗ', 1, 2, '#003366'),
        ('Волна 2: 3 ТЗ', 2, 3, '#0066CC'),
        ('Волна 3: 2 ТЗ', 3, 4, '#FF6600'),
    ]

    for name, start, end, color in waves:
        ax.barh(0, end-start, left=start, height=0.5, color=color, edgecolor='black')
        ax.text((start+end)/2, 0, name, ha='center', va='center', color='white', fontweight='bold', fontsize=9)

    ax.set_xlim(0, 4)
    ax.set_ylim(-0.5, 0.5)
    ax.set_xticks([0, 1, 2, 3, 4])
    ax.set_xticklabels(['1 кв.\n2026', '2 кв.\n2026', '3 кв.\n2026', '4 кв.\n2026', '1 кв.\n2027'])
    ax.set_yticks([])
    ax.set_title('Дорожная карта внедрения автоматизации', fontsize=11, fontweight='bold')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_summary_infographic(filename):
    """Инфографика: Ключевые результаты проекта"""
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')

    # Заголовок
    ax.text(5, 5.5, 'КЛЮЧЕВЫЕ РЕЗУЛЬТАТЫ ПРОЕКТА', ha='center', fontsize=14, fontweight='bold', color='#003366')

    # Блоки с данными
    blocks = [
        (1.5, 4, '198', 'сотрудников\nв компании', '#003366'),
        (5, 4, '64', 'точки\nавтоматизации', '#0066CC'),
        (8.5, 4, '13', 'технических\nзаданий', '#339933'),
        (1.5, 2, '1 272', 'ч/мес\nэкономия', '#FF6600'),
        (5, 2, '13,9', 'млн ₽/год\nэффект', '#CC0000'),
        (8.5, 2, '69', 'регламентов\nподготовлено', '#666666'),
    ]

    for x, y, number, label, color in blocks:
        circle = plt.Circle((x, y), 0.8, color=color, alpha=0.9)
        ax.add_patch(circle)
        ax.text(x, y+0.1, number, ha='center', va='center', fontsize=16, fontweight='bold', color='white')
        ax.text(x, y-1.1, label, ha='center', va='top', fontsize=9, color='#333333')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def add_charts_to_document():
    """Добавляет диаграммы в документ"""
    print("Создание диаграмм...")

    # Создаём диаграммы
    charts = {
        'summary': (create_summary_infographic, f"{TEMP_DIR}/summary.png"),
        'departments': (create_pie_chart_departments, f"{TEMP_DIR}/departments.png"),
        'automation_categories': (create_pie_chart_automation_categories, f"{TEMP_DIR}/auto_categories.png"),
        'automation_waves': (create_bar_chart_automation, f"{TEMP_DIR}/waves.png"),
        'tz_savings': (create_horizontal_bar_tz, f"{TEMP_DIR}/tz_savings.png"),
        'timeline': (create_timeline_chart, f"{TEMP_DIR}/timeline.png"),
    }

    for name, (func, path) in charts.items():
        print(f"  Создание {name}...")
        func(path)

    print("\nДобавление диаграмм в документ...")
    doc = Document("/Users/Rita/Documents/Vibecoding/14_Space Travel _отработка_Курсор/ФИНАЛЬНЫЙ_ОТЧЁТ_ЗАЩИТА_FORMATTED.docx")

    # Находим места для вставки диаграмм
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # После "РАЗДЕЛ 1: РЕЗЮМЕ" - вставляем инфографику
        if 'РЕЗЮМЕ ДЛЯ РУКОВОДСТВА' in text:
            # Вставляем после следующих нескольких параграфов (после таблицы ключевых результатов)
            for j in range(i+1, min(i+20, len(doc.paragraphs))):
                if 'Каскад данных' in doc.paragraphs[j].text:
                    p = doc.paragraphs[j]._element
                    # Создаём новый параграф с изображением
                    new_para = doc.add_paragraph()
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_para.add_run()
                    run.add_picture(f"{TEMP_DIR}/summary.png", width=Inches(6))
                    # Перемещаем его
                    p.addnext(new_para._element)
                    break

        # После "Распределение сотрудников по категориям" - диаграмма департаментов
        if 'Распределение сотрудников по категориям' in text:
            p = para._element
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run()
            run.add_picture(f"{TEMP_DIR}/departments.png", width=Inches(5))
            p.addnext(new_para._element)

        # После "Распределение по категориям" (автоматизация) - диаграмма категорий
        if text == '## 4.2. Распределение по категориям' or 'Распределение по категориям' in text and 'Рутина' not in text:
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if 'Интеграция' in doc.paragraphs[j].text and '13' in doc.paragraphs[j].text:
                    p = doc.paragraphs[j]._element
                    new_para = doc.add_paragraph()
                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_para.add_run()
                    run.add_picture(f"{TEMP_DIR}/auto_categories.png", width=Inches(4.5))
                    p.addnext(new_para._element)
                    break

        # После "План внедрения" - дорожная карта
        if 'План внедрения' in text or 'Дорожная карта' in text.lower():
            p = para._element
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run()
            run.add_picture(f"{TEMP_DIR}/timeline.png", width=Inches(6))
            p.addnext(new_para._element)

        # После "Распределение экономии по волнам" - диаграмма волн
        if 'Распределение экономии по волнам' in text:
            p = para._element
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = new_para.add_run()
            run.add_picture(f"{TEMP_DIR}/waves.png", width=Inches(5.5))
            p.addnext(new_para._element)

    # Сохраняем
    output_path = "/Users/Rita/Documents/Vibecoding/14_Space Travel _отработка_Курсор/ФИНАЛЬНЫЙ_ОТЧЁТ_ЗАЩИТА_FORMATTED.docx"
    doc.save(output_path)
    print(f"\nДокумент сохранён: {output_path}")

    # Удаляем временные файлы
    import shutil
    shutil.rmtree(TEMP_DIR, ignore_errors=True)
    print("Временные файлы удалены.")

if __name__ == "__main__":
    add_charts_to_document()
