#!/usr/bin/env python3
"""
Скрипт форматирования Word-документа для Space Travel.

Параметры форматирования:
- Поля 1 см со всех сторон
- Шрифт Tahoma, размер 9
- Заголовки: Tahoma 16, синий цвет
- Таблицы: на всю ширину (от края до края), с сеткой, шапка голубая
- Нумерация страниц внизу по центру
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Цвета
BLUE_HEADING = RGBColor(0, 51, 153)  # Синий для заголовков
LIGHT_BLUE_BG = "B8CCE4"  # Голубой для шапки таблиц

# Размеры (A4: 21 см x 29.7 см, поля 1 см)
PAGE_WIDTH_CM = 21
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 1
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM  # 19 см
CONTENT_WIDTH_TWIPS = int(CONTENT_WIDTH_CM * 567)  # 10773 twips


def set_narrow_margins(doc):
    """Установить поля 1 см и размер A4."""
    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)


def add_page_numbers(doc):
    """Добавить нумерацию страниц внизу по центру."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p.clear()

        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        run.font.name = 'Tahoma'
        run.font.size = Pt(9)


def set_font_style(run, font_name='Tahoma', font_size=9, color=None, bold=False):
    """Установить шрифт."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def is_heading(paragraph):
    """Проверить, является ли абзац заголовком."""
    style_name = paragraph.style.name.lower() if paragraph.style else ''
    if 'heading' in style_name or 'заголовок' in style_name:
        return True
    text = paragraph.text.strip()
    if text.startswith('#') or text.startswith('РАЗДЕЛ') or text.startswith('ПРИЛОЖЕНИЕ'):
        return True
    return False


def remove_extra_spaces(text):
    """Удалить лишние пробелы."""
    if not text:
        return text
    text = re.sub(r' +', ' ', text)
    return text.strip()


def format_paragraph(paragraph, font_name='Tahoma', font_size=9):
    """Форматировать абзац."""
    is_head = is_heading(paragraph)

    for run in paragraph.runs:
        if run.text:
            run.text = remove_extra_spaces(run.text)
        if is_head:
            set_font_style(run, font_name, 16, BLUE_HEADING, bold=True)
        else:
            set_font_style(run, font_name, font_size)

    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(3) if is_head else Pt(0)
    paragraph.paragraph_format.space_after = Pt(6) if is_head else Pt(2)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)


def set_cell_shading(cell, color_hex):
    """Установить цвет фона ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def remove_element(parent, tag):
    """Удалить элемент по тегу."""
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)


def set_table_full_width(table):
    """Растянуть таблицу на всю ширину страницы без пробелов."""
    tbl = table._tbl

    # Получить или создать tblPr
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 1. Убрать отступ таблицы слева
    remove_element(tblPr, 'w:tblInd')
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # 2. Установить ширину таблицы = 100% ширины контента
    remove_element(tblPr, 'w:tblW')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(CONTENT_WIDTH_TWIPS))
    tblPr.append(tblW)

    # 3. Убрать промежутки между ячейками
    remove_element(tblPr, 'w:tblCellSpacing')
    cellSpacing = OxmlElement('w:tblCellSpacing')
    cellSpacing.set(qn('w:w'), '0')
    cellSpacing.set(qn('w:type'), 'dxa')
    tblPr.append(cellSpacing)

    # 4. Минимальные отступы внутри ячеек
    remove_element(tblPr, 'w:tblCellMar')
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), '28')  # ~0.5 мм
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)
    tblPr.append(tblCellMar)

    # 5. Фиксированная раскладка
    remove_element(tblPr, 'w:tblLayout')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # 6. Выравнивание по левому краю
    remove_element(tblPr, 'w:jc')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    tblPr.append(jc)

    # 7. Установить ширину колонок через tblGrid
    num_cols = len(table.columns)
    if num_cols > 0:
        col_width = CONTENT_WIDTH_TWIPS // num_cols

        # Удалить старый tblGrid
        old_grid = tbl.find(qn('w:tblGrid'))
        if old_grid is not None:
            tbl.remove(old_grid)

        # Создать новый tblGrid
        tblGrid = OxmlElement('w:tblGrid')
        for _ in range(num_cols):
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(col_width))
            tblGrid.append(gridCol)

        # Вставить после tblPr
        tblPr_index = list(tbl).index(tblPr)
        tbl.insert(tblPr_index + 1, tblGrid)

        # Установить ширину каждой ячейки
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                remove_element(tcPr, 'w:tcW')
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_width))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.insert(0, tcW)


def set_table_borders(table):
    """Установить границы таблицы (сетку)."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    remove_element(tblPr, 'w:tblBorders')
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def format_table(table, font_name='Tahoma', font_size=9):
    """Форматировать таблицу."""
    # Растянуть на всю ширину
    set_table_full_width(table)

    # Установить границы
    set_table_borders(table)

    # Форматировать строки
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE_BG)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = remove_extra_spaces(run.text)
                    if row_idx == 0:
                        set_font_style(run, font_name, font_size, bold=True)
                    else:
                        set_font_style(run, font_name, font_size)

                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.right_indent = Pt(0)


def format_document(input_path, output_path=None, font_name='Tahoma', font_size=9):
    """Форматировать весь документ."""
    print(f"Открываю документ: {input_path}")
    doc = Document(input_path)

    print("Устанавливаю поля 1 см...")
    set_narrow_margins(doc)

    print("Добавляю нумерацию страниц...")
    add_page_numbers(doc)

    print(f"Форматирую текст: {font_name}, {font_size}pt (заголовки 16pt синие)...")
    for paragraph in doc.paragraphs:
        format_paragraph(paragraph, font_name, font_size)

    print(f"Форматирую таблицы ({len(doc.tables)} шт.)...")
    for table in doc.tables:
        format_table(table, font_name, font_size)

    if output_path is None:
        output_path = input_path

    print(f"Сохраняю: {output_path}")
    doc.save(output_path)
    print("Готово!")

    return output_path


def main():
    if len(sys.argv) < 2:
        print("Использование: python format_docx.py <файл.docx> [выходной_файл.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not Path(input_path).exists():
        print(f"Ошибка: файл не найден: {input_path}")
        sys.exit(1)

    format_document(input_path, output_path)


if __name__ == "__main__":
    main()
